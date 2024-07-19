from flask import Flask, render_template, request, send_file, jsonify
import os
from crewai import Agent, Task, Crew
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
from dotenv import load_dotenv
from langchain_openai import AzureChatOpenAI
from docx import Document
from io import BytesIO

load_dotenv()

# Initialize Flask app
app = Flask(__name__)

# LLM object and API Key
os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")

# Initialize Azure OpenAI client
llm = AzureChatOpenAI(
    api_key=os.getenv("AZURE_OPENAI_KEY"),
    api_version="2024-02-15-preview",
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    model="gpt35turbo16k",
    temperature=0.1,
    max_tokens=8000,
)

def generate_docx(result):
    doc = Document()
    doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Define Agents
diagnostician = Agent(
    role="Medical Diagnostician",
    goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis.",
    backstory="This agent specializes in diagnosing medical conditions based on patient-reported symptoms and medical history. It uses advanced algorithms and medical knowledge to identify potential health issues.",
    verbose=True,
    allow_delegation=False,
    tools=[SerperDevTool(), ScrapeWebsiteTool()],
    llm=llm
)

treatment_advisor = Agent(
    role="Treatment Advisor",
    goal="Recommend appropriate treatment plans based on the diagnosis provided by the Medical Diagnostician.",
    backstory="This agent specializes in creating treatment plans tailored to individual patient needs. It considers the diagnosis, patient history, and current best practices in medicine to recommend effective treatments.",
    verbose=True,
    allow_delegation=False,
    tools=[SerperDevTool(), ScrapeWebsiteTool()],
    llm=llm
)

# Define Tasks
diagnose_task = Task(
    description=(
        "1. Analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).\n"
        "2. Provide a preliminary diagnosis with possible conditions based on the provided information.\n"
        "3. Limit the diagnosis to the most likely conditions."
        "4. Limit the total output to 16,000 words."
    ),
    expected_output="A preliminary diagnosis with a list of possible conditions.",
    agent=diagnostician
)

treatment_task = Task(
    description=(
        "1. Based on the diagnosis, recommend appropriate treatment plans step by step.\n"
        "2. Consider the patient's medical history ({medical_history}) and current symptoms ({symptoms}).\n"
        "3. Provide detailed treatment recommendations, including medications, lifestyle changes, and follow-up care."
        "4. Limit the total output to 16,000 words."
    ),
    expected_output="A comprehensive treatment plan tailored to the patient's needs.",
    agent=treatment_advisor
)

# Create Crew
crew = Crew(
    agents=[diagnostician, treatment_advisor],
    tasks=[diagnose_task, treatment_task],
    verbose=2
)

@app.route('/')
def index():
    return render_template('app.html')

@app.route('/diagnose', methods=['POST'])
def diagnose():
    gender = request.form.get('gender')
    age = request.form.get('age')
    symptoms = request.form.get('symptoms')
    medical_history = request.form.get('medical_history')

    # Kickoff the Crew with inputs
    result = crew.kickoff(inputs={"symptoms": symptoms, "medical_history": medical_history})
    
    # Generate DOCX
    docx_file = generate_docx(result)
    
    # Prepare response
    response = {
        "result": result,
        "download_link": request.url_root + 'download'
    }
    
    # Save the DOCX file to a temporary location
    with open("temp_diagnosis_and_treatment_plan.docx", "wb") as f:
        f.write(docx_file.getbuffer())

    return jsonify(response)

@app.route('/download')
def download():
    return send_file("temp_diagnosis_and_treatment_plan.docx", as_attachment=True, attachment_filename="diagnosis_and_treatment_plan.docx")

if __name__ == '__main__':
    app.run(debug=True)
